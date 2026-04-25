"""
src/ingestion/document_loader.py
──────────────────────────────────
Loads documents from a directory.
Supports: .txt, .md, .pdf, .docx
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """A raw document loaded from disk."""
    page_content: str
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    """
    Loads all supported documents from a directory (recursive).

    Usage:
        loader = DocumentLoader("data/sample_docs")
        docs = loader.load()
    """

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

    def __init__(self, docs_dir: str):
        self.docs_dir = Path(docs_dir)
        if not self.docs_dir.exists():
            raise FileNotFoundError(f"Directory not found: {docs_dir}")

    def load(self) -> List[Document]:
        """Load all supported documents from the directory."""
        documents: List[Document] = []
        files = [
            p for p in self.docs_dir.rglob("*")
            if p.suffix.lower() in self.SUPPORTED_EXTENSIONS
        ]

        if not files:
            logger.warning("No supported documents found in %s", self.docs_dir)
            return documents

        for file_path in files:
            try:
                docs = self._load_file(file_path)
                documents.extend(docs)
                logger.info("Loaded %d pages from %s", len(docs), file_path.name)
            except Exception as e:
                logger.error("Failed to load %s: %s", file_path, e)

        logger.info("Total documents loaded: %d", len(documents))
        return documents

    def _load_file(self, path: Path) -> List[Document]:
        ext = path.suffix.lower()
        if ext in {".txt", ".md"}:
            return self._load_text(path)
        elif ext == ".pdf":
            return self._load_pdf(path)
        elif ext == ".docx":
            return self._load_docx(path)
        return []

    def _load_text(self, path: Path) -> List[Document]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [Document(
            page_content=text,
            metadata={"source": str(path), "file_type": path.suffix}
        )]

    def _load_pdf(self, path: Path) -> List[Document]:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("Install pypdf: pip install pypdf")

        reader = PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": str(path), "page": i + 1, "file_type": ".pdf"}
                ))
        return docs

    def _load_docx(self, path: Path) -> List[Document]:
        try:
            import docx
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        doc = docx.Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(
            page_content=text,
            metadata={"source": str(path), "file_type": ".docx"}
        )]
